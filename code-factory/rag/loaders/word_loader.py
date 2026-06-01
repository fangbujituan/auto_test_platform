"""Word (.docx) 文档加载器 / Word (.docx) document loader for the RAG pipeline.

从 Word 文档中提取内容，同时保留结构信息，
Extracts content from Word documents while preserving structural
如标题、段落、代码块和表格。
information such as headings, paragraphs, code blocks, and tables.

使用 python-docx 进行文档解析。
Uses python-docx for document parsing.

需求 / Requirements: 5.1, 5.2
"""

from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table as DocxTable

from rag.document_loader import FormatLoader
from src.core.exceptions import DocumentLoadError
from src.core.logging import get_logger
from src.core.models import DocumentFormat, DocumentUnit, LoadedDocument

logger = get_logger("rag.loaders.word_loader")

# Heading style name prefixes in python-docx
HEADING_STYLE_PREFIX = "Heading"

# Font names commonly used for code
CODE_FONT_NAMES = {
    "courier",
    "courier new",
    "consolas",
    "monaco",
    "menlo",
    "source code pro",
    "fira code",
    "jetbrains mono",
    "lucida console",
    "andale mono",
    "dejavu sans mono",
    "liberation mono",
}


def _is_code_paragraph(paragraph) -> bool:
    """Detect if a paragraph is likely a code block.

    Checks for monospace fonts or code-related style names.
    """
    style_name = (paragraph.style.name or "").lower()

    # Check style name for code indicators
    if any(indicator in style_name for indicator in ("code", "mono", "preformatted", "literal")):
        return True

    # Check if the paragraph uses a monospace font
    for run in paragraph.runs:
        if run.font and run.font.name:
            font_name = run.font.name.lower()
            if font_name in CODE_FONT_NAMES:
                return True

    return False


def _get_heading_level(paragraph) -> int:
    """Extract heading level from paragraph style.

    Returns the heading level (1-9) or 0 if not a heading.
    """
    style_name = paragraph.style.name or ""
    if style_name.startswith(HEADING_STYLE_PREFIX):
        # Extract level number from style name like "Heading 1", "Heading 2"
        level_str = style_name[len(HEADING_STYLE_PREFIX):].strip()
        try:
            return int(level_str)
        except ValueError:
            return 1  # Default to level 1 if can't parse
    return 0


def _table_to_text(table: DocxTable) -> str:
    """Convert a Word table to a pipe-delimited text representation.

    Preserves the table structure in a readable format.
    """
    rows_text: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows_text.append("| " + " | ".join(cells) + " |")

    # Add separator after header row if there are multiple rows
    if len(rows_text) > 1:
        header = rows_text[0]
        col_count = header.count("|") - 1
        separator = "|" + "|".join([" --- "] * col_count) + "|"
        rows_text.insert(1, separator)

    return "\n".join(rows_text)


class WordLoader(FormatLoader):
    """Word (.docx) 格式加载器，使用 python-docx / Word (.docx) format loader using python-docx.

    从 Word 文档中提取段落、标题、表格和代码块，
    Extracts paragraphs, headings, tables, and code blocks from
    保留结构信息。
    Word documents, preserving structural information.
    """

    def load(self, file_path: str, content: bytes) -> LoadedDocument:
        """Load a Word document and extract structured content.

        Args:
            file_path: Path to the source .docx file.
            content: Raw .docx file content as bytes.

        Returns:
            A LoadedDocument with extracted units and structural info.

        Raises:
            DocumentLoadError: If the document cannot be parsed.
        """
        import io

        try:
            doc = DocxDocument(io.BytesIO(content))
        except Exception as e:
            raise DocumentLoadError(
                f"Failed to parse Word document '{file_path}': {e}"
            )

        units: list[DocumentUnit] = []
        raw_text_parts: list[str] = []
        position = 0

        # Track which elements we've processed (tables can appear between paragraphs)
        # python-docx exposes document body elements in order
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # It's a paragraph element - find the matching paragraph object
                para = self._find_paragraph_for_element(doc, element)
                if para is None:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                raw_text_parts.append(text)

                # Check if it's a heading
                heading_level = _get_heading_level(para)
                if heading_level > 0:
                    units.append(DocumentUnit(
                        content=text,
                        unit_type="heading",
                        metadata={"heading_level": heading_level},
                        source_path=file_path,
                        position=position,
                    ))
                    position += 1
                    continue

                # Check if it's a code block
                if _is_code_paragraph(para):
                    # Collect consecutive code paragraphs
                    units.append(DocumentUnit(
                        content=text,
                        unit_type="code_block",
                        metadata={"style": para.style.name},
                        source_path=file_path,
                        position=position,
                    ))
                    position += 1
                    continue

                # Regular paragraph
                units.append(DocumentUnit(
                    content=text,
                    unit_type="paragraph",
                    metadata={},
                    source_path=file_path,
                    position=position,
                ))
                position += 1

            elif tag == "tbl":
                # It's a table element - find the matching table object
                table = self._find_table_for_element(doc, element)
                if table is None:
                    continue

                table_text = _table_to_text(table)
                if table_text.strip():
                    raw_text_parts.append(table_text)
                    row_count = len(table.rows)
                    col_count = len(table.columns)
                    units.append(DocumentUnit(
                        content=table_text,
                        unit_type="table",
                        metadata={
                            "row_count": row_count,
                            "col_count": col_count,
                        },
                        source_path=file_path,
                        position=position,
                    ))
                    position += 1

        # Merge consecutive code_block units that are adjacent
        units = self._merge_consecutive_code_blocks(units, file_path)

        raw_text = "\n\n".join(raw_text_parts)

        structural_info: dict[str, Any] = {
            "unit_count": len(units),
            "unit_types": list(set(u.unit_type for u in units)),
            "paragraph_count": sum(1 for u in units if u.unit_type == "paragraph"),
            "heading_count": sum(1 for u in units if u.unit_type == "heading"),
            "table_count": sum(1 for u in units if u.unit_type == "table"),
            "code_block_count": sum(1 for u in units if u.unit_type == "code_block"),
        }

        logger.info(
            "Word document loaded",
            file_path=file_path,
            unit_count=len(units),
        )

        return LoadedDocument(
            source_path=file_path,
            format=DocumentFormat.WORD,
            units=units,
            raw_text=raw_text,
            structural_info=structural_info,
        )

    def _find_paragraph_for_element(self, doc: DocxDocument, element) -> Any:
        """Find the python-docx Paragraph object corresponding to an XML element."""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table_for_element(self, doc: DocxDocument, element) -> Any:
        """Find the python-docx Table object corresponding to an XML element."""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    def _merge_consecutive_code_blocks(
        self, units: list[DocumentUnit], source_path: str
    ) -> list[DocumentUnit]:
        """Merge consecutive code_block units into single blocks.

        When multiple consecutive paragraphs are detected as code,
        they should be merged into a single code_block unit.
        """
        if not units:
            return units

        merged: list[DocumentUnit] = []
        i = 0
        position = 0

        while i < len(units):
            unit = units[i]

            if unit.unit_type == "code_block":
                # Collect consecutive code blocks
                code_parts = [unit.content]
                i += 1
                while i < len(units) and units[i].unit_type == "code_block":
                    code_parts.append(units[i].content)
                    i += 1

                merged.append(DocumentUnit(
                    content="\n".join(code_parts),
                    unit_type="code_block",
                    metadata={"line_count": len(code_parts)},
                    source_path=source_path,
                    position=position,
                ))
                position += 1
            else:
                merged.append(DocumentUnit(
                    content=unit.content,
                    unit_type=unit.unit_type,
                    metadata=unit.metadata,
                    source_path=source_path,
                    position=position,
                ))
                position += 1
                i += 1

        return merged
