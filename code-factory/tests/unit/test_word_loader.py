"""Unit tests for the Word (.docx) document loader.

Tests Word document parsing with structural preservation
(headings, paragraphs, code blocks, tables).

Requirements: 5.1, 5.2
"""

import io
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt

from rag.loaders.word_loader import (
    WordLoader,
    _get_heading_level,
    _is_code_paragraph,
    _table_to_text,
)
from src.core.exceptions import DocumentLoadError
from src.core.models import DocumentFormat, DocumentUnit, LoadedDocument


def _create_simple_docx(paragraphs: list[tuple[str, str]]) -> bytes:
    """Create a simple .docx file with given paragraphs.

    Args:
        paragraphs: List of (text, style_name) tuples.
            style_name can be "Normal", "Heading 1", "Heading 2", etc.

    Returns:
        Bytes content of the .docx file.
    """
    doc = DocxDocument()
    for text, style in paragraphs:
        doc.add_paragraph(text, style=style)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_docx_with_table(headers: list[str], rows: list[list[str]]) -> bytes:
    """Create a .docx file with a table.

    Args:
        headers: Column headers.
        rows: List of row data (each row is a list of cell values).

    Returns:
        Bytes content of the .docx file.
    """
    doc = DocxDocument()
    doc.add_paragraph("Table Section", style="Heading 1")
    table = doc.add_table(rows=1, cols=len(headers))
    # Add headers
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    # Add data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _create_docx_with_code(code_text: str) -> bytes:
    """Create a .docx file with a code-styled paragraph.

    Uses a monospace font to simulate code blocks.
    """
    doc = DocxDocument()
    doc.add_paragraph("Code Example", style="Heading 1")
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestWordLoader:
    """Tests for the WordLoader class."""

    def test_load_simple_document(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("Introduction", "Heading 1"),
            ("This is the first paragraph.", "Normal"),
            ("Details", "Heading 2"),
            ("More detailed content here.", "Normal"),
        ])
        result = loader.load("test.docx", content)

        assert isinstance(result, LoadedDocument)
        assert result.source_path == "test.docx"
        assert result.format == DocumentFormat.WORD
        assert len(result.units) == 4

    def test_load_extracts_headings(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("Main Title", "Heading 1"),
            ("Some text.", "Normal"),
            ("Sub Section", "Heading 2"),
        ])
        result = loader.load("test.docx", content)

        headings = [u for u in result.units if u.unit_type == "heading"]
        assert len(headings) == 2
        assert headings[0].content == "Main Title"
        assert headings[0].metadata["heading_level"] == 1
        assert headings[1].content == "Sub Section"
        assert headings[1].metadata["heading_level"] == 2

    def test_load_extracts_paragraphs(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("First paragraph content.", "Normal"),
            ("Second paragraph content.", "Normal"),
        ])
        result = loader.load("test.docx", content)

        paragraphs = [u for u in result.units if u.unit_type == "paragraph"]
        assert len(paragraphs) == 2
        assert paragraphs[0].content == "First paragraph content."
        assert paragraphs[1].content == "Second paragraph content."

    def test_load_extracts_tables(self):
        loader = WordLoader()
        content = _create_docx_with_table(
            headers=["Name", "Age"],
            rows=[["Alice", "30"], ["Bob", "25"]],
        )
        result = loader.load("test.docx", content)

        tables = [u for u in result.units if u.unit_type == "table"]
        assert len(tables) == 1
        assert "Name" in tables[0].content
        assert "Alice" in tables[0].content
        assert tables[0].metadata["row_count"] == 3  # header + 2 data rows
        assert tables[0].metadata["col_count"] == 2

    def test_load_extracts_code_blocks(self):
        loader = WordLoader()
        content = _create_docx_with_code("def hello():\n    print('world')")
        result = loader.load("test.docx", content)

        code_blocks = [u for u in result.units if u.unit_type == "code_block"]
        assert len(code_blocks) >= 1
        assert "hello" in code_blocks[0].content

    def test_load_invalid_content_raises_error(self):
        loader = WordLoader()
        with pytest.raises(DocumentLoadError) as exc_info:
            loader.load("bad.docx", b"this is not a docx file")
        assert "Failed to parse Word document" in str(exc_info.value)

    def test_load_empty_document(self):
        loader = WordLoader()
        doc = DocxDocument()
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        result = loader.load("empty.docx", content)
        assert isinstance(result, LoadedDocument)
        assert result.units == []

    def test_units_have_sequential_positions(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("Title", "Heading 1"),
            ("Para 1.", "Normal"),
            ("Para 2.", "Normal"),
            ("Sub Title", "Heading 2"),
            ("Para 3.", "Normal"),
        ])
        result = loader.load("test.docx", content)

        positions = [u.position for u in result.units]
        assert positions == list(range(len(positions)))

    def test_units_have_correct_source_path(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("Hello", "Heading 1"),
            ("World.", "Normal"),
        ])
        result = loader.load("/docs/report.docx", content)

        for unit in result.units:
            assert unit.source_path == "/docs/report.docx"

    def test_structural_info_populated(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("Title", "Heading 1"),
            ("Content.", "Normal"),
        ])
        result = loader.load("test.docx", content)

        assert "unit_count" in result.structural_info
        assert "heading_count" in result.structural_info
        assert "paragraph_count" in result.structural_info
        assert result.structural_info["unit_count"] == len(result.units)

    def test_raw_text_contains_content(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("Title", "Heading 1"),
            ("Some important text.", "Normal"),
        ])
        result = loader.load("test.docx", content)

        assert "Title" in result.raw_text
        assert "Some important text" in result.raw_text

    def test_units_have_valid_types(self):
        loader = WordLoader()
        content = _create_simple_docx([
            ("Heading", "Heading 1"),
            ("Paragraph.", "Normal"),
        ])
        result = loader.load("test.docx", content)

        valid_types = {"heading", "paragraph", "code_block", "table"}
        for unit in result.units:
            assert unit.unit_type in valid_types


class TestTableToText:
    """Tests for the _table_to_text helper function."""

    def test_simple_table(self):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"

        text = _table_to_text(table)
        assert "| A | B |" in text
        assert "| 1 | 2 |" in text
        assert "---" in text  # separator

    def test_single_row_table(self):
        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "X"
        table.rows[0].cells[1].text = "Y"
        table.rows[0].cells[2].text = "Z"

        text = _table_to_text(table)
        assert "| X | Y | Z |" in text
        # No separator for single row
        assert "---" not in text


class TestGetHeadingLevel:
    """Tests for the _get_heading_level helper function."""

    def test_heading_1(self):
        doc = DocxDocument()
        para = doc.add_paragraph("Title", style="Heading 1")
        assert _get_heading_level(para) == 1

    def test_heading_2(self):
        doc = DocxDocument()
        para = doc.add_paragraph("Subtitle", style="Heading 2")
        assert _get_heading_level(para) == 2

    def test_heading_3(self):
        doc = DocxDocument()
        para = doc.add_paragraph("Sub-subtitle", style="Heading 3")
        assert _get_heading_level(para) == 3

    def test_normal_paragraph_returns_0(self):
        doc = DocxDocument()
        para = doc.add_paragraph("Normal text", style="Normal")
        assert _get_heading_level(para) == 0


class TestIsCodeParagraph:
    """Tests for the _is_code_paragraph helper function."""

    def test_monospace_font_detected(self):
        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("code here")
        run.font.name = "Courier New"
        assert _is_code_paragraph(para) is True

    def test_consolas_font_detected(self):
        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("more code")
        run.font.name = "Consolas"
        assert _is_code_paragraph(para) is True

    def test_normal_font_not_code(self):
        doc = DocxDocument()
        para = doc.add_paragraph("Normal text", style="Normal")
        assert _is_code_paragraph(para) is False
